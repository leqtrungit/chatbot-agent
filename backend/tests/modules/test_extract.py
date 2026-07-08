from __future__ import annotations

import pytest
from docx import Document as DocxDocument

from app.modules.document.pipeline.extract import extract_text


def test_extract_txt(tmp_path):
    path = tmp_path / "a.txt"
    path.write_text("hello plain text", encoding="utf-8")
    assert extract_text(str(path), "text/plain") == "hello plain text"


def test_extract_md(tmp_path):
    path = tmp_path / "a.md"
    path.write_text("# heading\nbody", encoding="utf-8")
    assert extract_text(str(path), "text/markdown") == "# heading\nbody"


def test_extract_docx(tmp_path):
    path = tmp_path / "a.docx"
    doc = DocxDocument()
    doc.add_paragraph("first paragraph")
    doc.add_paragraph("second paragraph")
    doc.save(str(path))

    text = extract_text(
        str(path),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "first paragraph" in text
    assert "second paragraph" in text


def test_extract_pdf(tmp_path):
    pypdf = pytest.importorskip("pypdf")
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    path = tmp_path / "a.pdf"
    with open(path, "wb") as f:
        writer.write(f)

    # A blank page has no text; just verify the PDF branch runs without error.
    text = extract_text(str(path), "application/pdf")
    assert isinstance(text, str)


def test_extract_unsupported_mime_raises(tmp_path):
    path = tmp_path / "a.bin"
    path.write_bytes(b"\x00\x01")
    with pytest.raises(ValueError):
        extract_text(str(path), "application/octet-stream")
